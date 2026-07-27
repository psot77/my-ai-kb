import os
import socket
import json
import urllib.request
import logging
import io
import asyncio
import time
import uuid
import datetime
import requests

# Чтение PDF и DOCX
import pypdf
import docx

# Настройка логирования
logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)

# =====================================================================
# 1. КОНСТАНТЫ И ПАРАМЕТРЫ ИНДЕКСАЦИИ
# =====================================================================
CHUNK_SIZE = 500
CHUNK_OVERLAP = 100
EMBEDDING_MODEL = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
VECTOR_DIM = 384
QDRANT_URL = "https://18545c10-4b80-4ed2-9304-4ba636a29618.eu-west-1-0.aws.cloud.qdrant.io"
COLLECTION_NAME = "knowledge_base"
ANALYTICS_COLLECTION = "analytics_logs"

# =====================================================================
# 2. DNS-OVER-HTTPS (DoH)
# =====================================================================
DNS_CACHE = {}

def resolve_via_google_doh(hostname: str) -> str:
    if hostname in DNS_CACHE:
        return DNS_CACHE[hostname]
    try:
        url = f"https://8.8.8.8/resolve?name={hostname}&type=A"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=5) as response:
            data = json.loads(response.read().decode())
            if data.get("Status") == 0 and "Answer" in data:
                for ans in data["Answer"]:
                    if ans.get("type") == 1:
                        ip = ans.get("data")
                        DNS_CACHE[hostname] = ip
                        return ip
    except Exception as e:
        logger.warning(f"DoH error: {e}")
    return None

old_getaddrinfo = socket.getaddrinfo

def custom_getaddrinfo(host, port, family=0, type=0, proto=0, flags=0):
    try:
        return old_getaddrinfo(host, port, socket.AF_INET, type, proto, flags)
    except socket.gaierror:
        ip = resolve_via_google_doh(host)
        if ip:
            return old_getaddrinfo(ip, port, socket.AF_INET, type, proto, flags)
        raise

socket.getaddrinfo = custom_getaddrinfo

# =====================================================================
# 3. НАСТРОЙКИ КЛИЕНТОВ И АДМИНОВ
# =====================================================================
from telegram import Update
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters, ContextTypes
from qdrant_client import QdrantClient
from qdrant_client.models import PointStruct, VectorParams, Distance, Filter, FieldCondition, MatchValue
from groq import Groq
from huggingface_hub import InferenceClient

logger.info("=== СТАРТ BOT.PY (С ПРИВЯЗКОЙ К ОБЩЕМУ РАЗДЕЛУ) ===")

def clean_env(var_name: str, fallback: str = None) -> str:
    val = os.getenv(var_name, fallback)
    if val:
        val = val.strip().strip("'").strip('"')
    return val

TELEGRAM_BOT_TOKEN = clean_env("TELEGRAM_BOT_TOKEN", "8901191309:AAF4UuKO5RIZX7_Z2mj7PKp7K-chKZJdvE8")
GROQ_API_KEY = clean_env("GROQ_API_KEY")
QDRANT_API_KEY = clean_env("QDRANT_API_KEY")
HF_TOKEN = clean_env("HF_TOKEN")

ADMIN_IDS_RAW = clean_env("ADMIN_IDS", "")
ADMIN_IDS = {int(x.strip()) for x in ADMIN_IDS_RAW.split(",") if x.strip().isdigit()}

qdrant = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, port=443, https=True, check_compatibility=False)
groq_client = Groq(api_key=GROQ_API_KEY)

def is_admin(user_id: int) -> bool:
    if not ADMIN_IDS:
        return True
    return user_id in ADMIN_IDS

try:
    cols = [c.name for c in qdrant.get_collections().collections]
    if ANALYTICS_COLLECTION not in cols:
        qdrant.create_collection(
            collection_name=ANALYTICS_COLLECTION,
            vectors_config=VectorParams(size=1, distance=Distance.COSINE)
        )
        logger.info(f"Создана коллекция аналитики '{ANALYTICS_COLLECTION}'")
except Exception as e:
    logger.warning(f"Ошибка проверки коллекции аналитики: {e}")

# =====================================================================
# 4. ФУНКЦИЯ ЗАПИСИ АНАЛИТИКИ
# =====================================================================
def log_analytics(source: str, user_id: str, username: str, event_type: str, query: str = "", score: float = 0.0, status: str = "Success", details: str = ""):
    try:
        now_str = datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%d %H:%M:%S")
        point = PointStruct(
            id=str(uuid.uuid4()),
            vector=[0.0],
            payload={
                "timestamp": now_str,
                "source": source,
                "user_id": str(user_id),
                "username": username or "Аноним",
                "event_type": event_type,
                "query": query[:300],
                "score": float(score),
                "found_in_kb": bool(score >= 0.20),
                "status": status,
                "details": details
            }
        )
        qdrant.upsert(collection_name=ANALYTICS_COLLECTION, points=[point])
    except Exception as e:
        logger.error(f"Ошибка записи аналитики в Qdrant: {e}")

# =====================================================================
# 5. ВЕКТОРИЗАЦИЯ И ИЗВЛЕЧЕНИЕ ТЕКСТА
# =====================================================================
def get_cloud_embedding(text: str) -> list:
    client = InferenceClient(token=HF_TOKEN)
    last_error = None
    for attempt in range(3):
        try:
            result = client.feature_extraction(
                text,
                model=EMBEDDING_MODEL
            )
            data = result.tolist() if hasattr(result, "tolist") else result
            while isinstance(data, list) and len(data) > 0 and isinstance(data[0], list):
                data = data[0]
            if isinstance(data, list) and len(data) > 0 and isinstance(data[0], (int, float)):
                return [float(x) for x in data]
            raise Exception(f"Неожиданный формат HF: {data}")
        except Exception as e:
            last_error = str(e)
            time.sleep(2)
    raise Exception(f"Ошибка вектора: {last_error}")

def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    doc_obj = docx.Document(io.BytesIO(file_bytes))
    full_text = []

    for p in doc_obj.paragraphs:
        if p.text.strip():
            full_text.append(p.text.strip())

    for table in doc_obj.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))

    return "\n".join(full_text)

def split_text_into_chunks(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    if not text:
        return []
    chunks = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = start + chunk_size
        chunk = text[start:end]
        if end < text_len and not text[end].isspace():
            last_space = chunk.rfind(' ')
            if last_space != -1:
                end = start + last_space
                chunk = text[start:end]
        chunks.append(chunk.strip())
        start += (chunk_size - overlap)
    return [c for c in chunks if len(c) > 20]

# =====================================================================
# 6. КОМАНДЫ И ОБРАБОТЧИКИ TELEGRAM
# =====================================================================
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 **Здравствуйте! Я ваш виртуальный корпоративный ассистент.**\n\n"
        "💬 Вы можете задать мне любой вопрос по базе знаний.\n\n"
        "🛠️ **Полезные команды:**\n"
        "• `/id` — Узнать ваш цифровой Telegram ID\n"
        "• `/stats` — Статистика обращения к базе\n"
        "• `/files` — Список документов в базе\n"
        "• `/delete имя_файла.docx` — Удалить файл из базы",
        parse_mode="Markdown"
    )

async def get_my_id(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    username = update.effective_user.username or update.effective_user.first_name
    admin_status = "👑 Администратор" if is_admin(user_id) else "👤 Пользователь"
    
    await update.message.reply_text(
        f"🆔 **Ваши данные:**\n"
        f"• Telegram ID: `{user_id}`\n"
        f"• Пользователь: @{username}\n"
        f"• Статус в боте: **{admin_status}**",
        parse_mode="Markdown"
    )

async def stats(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if not is_admin(user_id):
        await update.message.reply_text("⛔ **У вас нет прав для просмотра статистики.**", parse_mode="Markdown")
        return

    try:
        kb_count_res = qdrant.count(collection_name=COLLECTION_NAME)
        total_chunks = kb_count_res.count

        analytics_res, _ = qdrant.scroll(
            collection_name=ANALYTICS_COLLECTION,
            limit=1000,
            with_payload=True,
            with_vectors=False
        )

        total_requests = len(analytics_res) if analytics_res else 0
        tg_requests = 0
        web_requests = 0
        successful_requests = 0

        if analytics_res:
            for pt in analytics_res:
                p = pt.payload or {}
                if p.get("source") == "Telegram":
                    tg_requests += 1
                elif p.get("source") == "Web":
                    web_requests += 1
                if p.get("found_in_kb"):
                    successful_requests += 1

        success_rate = round((successful_requests / total_requests) * 100, 1) if total_requests > 0 else 0.0

        msg_text = (
            f"📊 **СТАТИСТИКА AI АССИСТЕНТА**\n\n"
            f"🧠 **База знаний Qdrant:**\n"
            f"• Всего чанков в базе: `{total_chunks}` шт.\n\n"
            f"📈 **Аналитика обращений:**\n"
            f"• Всего запросов: `{total_requests}`\n"
            f"• Из Telegram 📱: `{tg_requests}`\n"
            f"• Из Веб-панели 🌐: `{web_requests}`\n"
            f"• Успешность ответов (RAG): `{success_rate}%`\n"
        )
        await update.message.reply_text(msg_text, parse_mode="Markdown")
    except Exception as e:
        logger.error(f"Ошибка команды /stats: {e}", exc_info=True)
        await update.message.reply_text(f"⚠️ Ошибка получения статистики: {e}")

async def files_list(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if not is_admin(user_id):
        await update.message.reply_text("⛔ **У вас нет прав для просмотра списка файлов.**", parse_mode="Markdown")
        return

    try:
        scroll_res, _ = qdrant.scroll(
            collection_name=COLLECTION_NAME,
            limit=10000,
            with_payload=["source_file"],
            with_vectors=False
        )

        files_summary = {}
        for pt in scroll_res:
            p = pt.payload or {}
            src = p.get("source_file", "Неизвестный файл")
            files_summary[src] = files_summary.get(src, 0) + 1

        if not files_summary:
            await update.message.reply_text("📂 В базе знаний пока нет загруженных файлов.")
            return

        lines = ["📂 **Список файлов в базе знаний:**\n"]
        for idx, (fname, count) in enumerate(files_summary.items(), start=1):
            lines.append(f"{idx}. `{fname}` — *{count} чанков*")

        lines.append("\n💡 *Чтобы удалить файл из базы знаний, используйте команду:*")
        lines.append("`/delete точное_имя_файла.docx`")

        await update.message.reply_text("\n".join(lines), parse_mode="Markdown")
    except Exception as e:
        logger.error(f"Ошибка команды /files: {e}", exc_info=True)
        await update.message.reply_text(f"⚠️ Ошибка получения списка файлов: {e}")

async def delete_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    if not is_admin(user.id):
        await update.message.reply_text("⛔ **У вас нет прав для удаления файлов.**", parse_mode="Markdown")
        return

    if not context.args:
        await update.message.reply_text(
            "⚠️ **Укажите имя файла для удаления.**\n\n"
            "Пример: `/delete КОМЕРЦІЙНА ПРОПОЗИЦІЯ.docx`\n\n"
            "📋 Посмотреть список всех файлов можно по команде `/files`",
            parse_mode="Markdown"
        )
        return

    target_filename = " ".join(context.args).strip()

    try:
        pts, _ = qdrant.scroll(
            collection_name=COLLECTION_NAME,
            scroll_filter=Filter(
                must=[FieldCondition(key="source_file", match=MatchValue(value=target_filename))]
            ),
            limit=10000,
            with_payload=False,
            with_vectors=False
        )

        p_ids = [p.id for p in pts]

        if not p_ids:
            await update.message.reply_text(
                f"❌ Файл `{target_filename}` не найден в базе знаний.\n\nПроверьте точное имя файла с помощью команды `/files`",
                parse_mode="Markdown"
            )
            return

        qdrant.delete(collection_name=COLLECTION_NAME, points_selector=p_ids)

        log_analytics(
            source="Telegram",
            user_id=user.id,
            username=user.username or user.first_name,
            event_type="Удаление файла",
            query=f"Удаление: {target_filename}",
            score=1.0,
            status="Удалено",
            details=f"Удалено чанков: {len(p_ids)}"
        )

        await update.message.reply_text(
            f"✅ **Файл успешно удален из базы знаний!**\n\n"
            f"📄 Имя файла: `{target_filename}`\n"
            f"🗑️ Удалено чанков из Qdrant: `{len(p_ids)}`",
            parse_mode="Markdown"
        )
    except Exception as e:
        logger.error(f"Ошибка при удалении файла: {e}", exc_info=True)
        await update.message.reply_text(f"⚠️ Ошибка при удалении файла: {e}")

def search_rag_answer(query_text: str, user_info: dict) -> str:
    try:
        query_vector = get_cloud_embedding(query_text)
        response = qdrant.query_points(
            collection_name=COLLECTION_NAME,
            query=query_vector,
            limit=5
        )
        search_results = response.points
        max_score = max([hit.score for hit in search_results]) if search_results else 0.0

        log_analytics(
            source="Telegram",
            user_id=user_info.get("id"),
            username=user_info.get("username"),
            event_type=user_info.get("event_type", "Текстовый запрос"),
            query=query_text,
            score=max_score,
            status="Успешно" if max_score >= 0.20 else "Не найдено в БЗ"
        )

        if not search_results or max_score < 0.20:
            return "К сожалению, в корпоративной базе знаний пока нет подробных инструкций по этому вопросу."

        context_chunks = [
            f"[Файл: {hit.payload.get('source_file', 'Документ')}]\n{hit.payload.get('text', '')}"
            for hit in search_results
        ]
        context = "\n\n---\n\n".join(context_chunks)

        llm_prompt = f"""Ты — высококвалифицированный корпоративный AI-ассистент базы знаний. 
Твоя задача — давать точные, профессиональные и структурированные ответы.

--- ПРАВИЛА И ОГРАНИЧЕНИЯ ---
1. **Язык ответа:** Отвечай СТРОГО на том же языке, на котором написан «ВОПРОС ПОЛЬЗОВАТЕЛЯ».
2. **Строгая точность (БЕЗ ГАЛЛЮЦИНАЦИЙ):** Используй ТОЛЬКО информацию из блока «ИНФОРМАЦИЯ ИЗ БАЗЫ ЗНАНИЙ». Не придумывай факты или цены. Если ответа нет, вежливо укажи, что информации в БЗ недостаточно.
3. **Форматирование:** Используй списки (`•` или `1.`), выделяй ключевые термины и цифры **жирным шрифтом**. Избегай «воды».
4. **Указание источников:** В конце ответа укажи названия файлов-источников, если они есть.

--- ИНФОРМАЦИЯ ИЗ БАЗЫ ЗНАНИЙ ---
{context}

--- ВОПРОС ПОЛЬЗОВАТЕЛЯ ---
{query_text}

--- ОТВЕТ ---"""

        res = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": llm_prompt}],
            temperature=0.1
        )
        return res.choices[0].message.content
    except Exception as e:
        logger.error(f"Ошибка RAG: {e}", exc_info=True)
        log_analytics(
            source="Telegram",
            user_id=user_info.get("id"),
            username=user_info.get("username"),
            event_type=user_info.get("event_type", "Текстовый запрос"),
            query=query_text,
            status="Ошибка",
            details=str(e)
        )
        return f"⚠️ Произошла ошибка при поиске: {e}"

async def handle_text_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        user_text = update.message.text
        user_info = {
            "id": update.effective_user.id,
            "username": update.effective_user.username or update.effective_user.first_name,
            "event_type": "Текстовый запрос"
        }
        await update.message.reply_chat_action("typing")
        loop = asyncio.get_running_loop()
        answer = await loop.run_in_executor(None, search_rag_answer, user_text, user_info)
        await update.message.reply_text(answer)
    except Exception as e:
        await update.message.reply_text(f"⚠️ Ошибка: {e}")

async def handle_voice_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        await update.message.reply_chat_action("typing")
        voice_file = await context.bot.get_file(update.message.voice.file_id)
        voice_bytes = await voice_file.download_as_bytearray()
        audio_io = io.BytesIO(voice_bytes)
        audio_io.name = "voice.ogg"

        transcription = groq_client.audio.transcriptions.create(
            file=audio_io,
            model="whisper-large-v3-turbo",
            prompt="Запрос на русском языке по базе знаний",
            response_format="text"
        )
        prompt_text = str(transcription).strip()
        await update.message.reply_text(f"🎙️ *Распознано:* `{prompt_text}`", parse_mode="Markdown")
        
        user_info = {
            "id": update.effective_user.id,
            "username": update.effective_user.username or update.effective_user.first_name,
            "event_type": "Голосовой запрос"
        }
        loop = asyncio.get_running_loop()
        answer = await loop.run_in_executor(None, search_rag_answer, prompt_text, user_info)
        await update.message.reply_text(answer)
    except Exception as e:
        await update.message.reply_text(f"⚠️ Ошибка при обработке аудио: {e}")

async def handle_document_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    if not is_admin(user.id):
        await update.message.reply_text("⛔ **У вас нет прав для добавления документов в базу знаний.**", parse_mode="Markdown")
        return

    doc = update.message.document
    file_name = doc.file_name
    file_size_kb = round(doc.file_size / 1024, 1)
    ext = os.path.splitext(file_name)[1].lower()

    if ext not in [".docx", ".pdf", ".txt"]:
        await update.message.reply_text("⚠️ Поддерживаются только форматы `.docx`, `.pdf` и `.txt`", parse_mode="Markdown")
        return

    start_time = time.time()

    msg = await update.message.reply_text(
        f"⚙️ **Запуск индексации файла...**\n\n"
        f"📁 **Файл:** `{file_name}` ({file_size_kb} KB)\n"
        f"⏳ *Идет чтение файла и генерация векторов...*",
        parse_mode="Markdown"
    )

    try:
        telegram_file = await context.bot.get_file(doc.file_id)
        file_bytes = await telegram_file.download_as_bytearray()
        
        extracted_text = ""
        if ext == ".txt":
            extracted_text = file_bytes.decode("utf-8", errors="ignore")
        elif ext == ".docx":
            extracted_text = extract_text_from_docx_bytes(file_bytes)
        elif ext == ".pdf":
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages = [page.extract_text() for page in reader.pages if page.extract_text()]
            extracted_text = "\n".join(pages)

        if not extracted_text.strip():
            await msg.edit_text(f"⚠️ Не удалось извлечь текст из файла `{file_name}`.")
            return

        chunks = split_text_into_chunks(extracted_text)
        
        points = []
        for idx, chunk in enumerate(chunks):
            vector = get_cloud_embedding(chunk)
            points.append(
                PointStruct(
                    id=str(uuid.uuid4()),
                    vector=vector,
                    payload={
                        "text": chunk,
                        "source_file": file_name,
                        "section": "Общий раздел",  # Привязка к разделу по умолчанию
                        "chunk_index": idx
                    }
                )
            )
            await asyncio.sleep(0.04)

        qdrant.upsert(collection_name=COLLECTION_NAME, points=points)
        elapsed_time = round(time.time() - start_time, 2)

        log_analytics(
            source="Telegram",
            user_id=user.id,
            username=user.username or user.first_name,
            event_type="Загрузка документа",
            query=f"Файл: {file_name}",
            score=1.0,
            status="Загружено",
            details=f"Чанков: {len(chunks)}, Время: {elapsed_time}s"
        )

        await msg.edit_text(
            f"✅ **Данные успешно проиндексированы и загружены!**\n\n"
            f"📊 **Статистика:**\n"
            f"• **Файл:** `{file_name}` (`{file_size_kb} KB`)\n"
            f"• **Раздел:** `Общий раздел`\n"
            f"• **Сгенерировано чанков:** `{len(chunks)}` шт.\n"
            f"• **Время обработки:** `{elapsed_time} сек`\n\n"
            f"💡 *Теперь можете задавать вопросы по содержанию этого документа!*",
            parse_mode="Markdown"
        )
    except Exception as e:
        logger.error(f"Ошибка загрузки файла: {e}", exc_info=True)
        await msg.edit_text(f"⚠️ Ошибка при индексации файла: {e}")

# =====================================================================
# 7. ЗАПУСК БОТА И РЕГИСТРАЦИЯ ХЕНДЛЕРОВ
# =====================================================================
if __name__ == '__main__':
    if not TELEGRAM_BOT_TOKEN:
        logger.critical("TELEGRAM_BOT_TOKEN не задан!")
        exit(1)
        
    app = ApplicationBuilder().token(TELEGRAM_BOT_TOKEN).build()
    
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("id", get_my_id))
    app.add_handler(CommandHandler("stats", stats))
    app.add_handler(CommandHandler("files", files_list))
    app.add_handler(CommandHandler("delete", delete_file))
    
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document_message))
    app.add_handler(MessageHandler(filters.TEXT & (~filters.COMMAND), handle_text_message))
    app.add_handler(MessageHandler(filters.VOICE, handle_voice_message))
    
    logger.info("🤖 УСПЕХ: Бот запущен!")
    app.run_polling(drop_pending_updates=True)
