import hashlib
import io
import json
import os
import time
import urllib.request
import uuid
from datetime import datetime

from docx import Document
from groq import Groq
from huggingface_hub import InferenceClient
from langchain_text_splitters import MarkdownHeaderTextSplitter
import pandas as pd
from pypdf import PdfReader
from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance,
    FieldCondition,
    Filter,
    MatchAny,
    MatchValue,
    PayloadSchemaType,
    PointStruct,
    VectorParams,
)
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate
import streamlit as st

# =====================================================================
# 1. НАСТРОЙКИ КЛЮЧЕЙ, СТРАНИЦЫ И ТАЙМ-АУТА
# =====================================================================
GROQ_API_KEY = (
    str(st.secrets.get("GROQ_API_KEY", "")).strip().strip("'").strip('"')
)
QDRANT_API_KEY = (
    str(st.secrets.get("QDRANT_API_KEY", "")).strip().strip("'").strip('"')
)
HF_TOKEN = str(st.secrets.get("HF_TOKEN", "")).strip().strip("'").strip('"')

QDRANT_URL = (
    "https://18545c10-4b80-4ed2-9304-4ba636a29618.eu-west-1-0.aws.cloud.qdrant.io"
)
COLLECTION_NAME = "knowledge_base"
LOGS_COLLECTION = "audit_logs"
ANALYTICS_COLLECTION = "analytics_logs"
CONFIG_COLLECTION = "system_config"
CHAT_HISTORY_COLLECTION = "chat_history"
RE_COLLECTION_NAME = "real_estate_listings"

SESSION_TIMEOUT_MINUTES = 15

st.set_page_config(
    page_title="Mavbot AI Enterprise",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =====================================================================
# CSS СТИЛИЗАЦИЯ (MAVBOT UI)
# =====================================================================
st.markdown(
    """
    <style>
    .stApp { background-color: #F8Fafd; }
    section[data-testid="stSidebar"] { background-color: #F0F4F9 !important; border-right: 1px solid #E1E8ED; padding-top: 10px; }
    .mavbot-header { display: flex; align-items: center; gap: 10px; font-size: 22px; font-weight: 700; color: #1F1F1F; margin-bottom: 15px; }
    .mavbot-icon { font-size: 26px; }
    .recent-title { font-size: 12px; font-weight: 600; color: #72777A; margin-top: 15px; margin-bottom: 8px; text-transform: uppercase; letter-spacing: 0.5px; }
    div[data-testid="stSidebar"] button { border-radius: 12px !important; border: none !important; text-align: left !important; font-size: 14px !important; color: #2D3135 !important; background-color: transparent !important; transition: all 0.2s ease; }
    div[data-testid="stSidebar"] button:hover { background-color: #E2E7EC !important; color: #1F1F1F !important; }
    div[data-testid="stSidebar"] div[data-testid="stPopover"] button { padding: 2px 6px !important; border-radius: 8px !important; font-size: 16px !important; }
    .user-profile-card { display: flex; align-items: center; gap: 10px; padding: 10px; background: #E8EEF5; border-radius: 16px; margin-top: 10px; }
    .user-avatar { width: 36px; height: 36px; border-radius: 50%; background-color: #4285F4; color: white; display: flex; align-items: center; justify-content: center; font-weight: bold; }
    .re-card { background-color: #FFFFFF; border: 1px solid #E1E8ED; border-radius: 14px; padding: 16px; margin-bottom: 12px; box-shadow: 0 2px 6px rgba(0,0,0,0.02); }
    .re-price { font-size: 20px; font-weight: 700; color: #1a73e8; }
    .re-badge { background-color: #e8f0fe; color: #1967d2; padding: 3px 8px; border-radius: 6px; font-size: 12px; font-weight: 600; margin-left: 4px; }
    .re-badge-owner { background-color: #e6f4ea; color: #137333; padding: 3px 8px; border-radius: 6px; font-size: 12px; font-weight: 600; margin-left: 4px; }
    ::-webkit-scrollbar { width: 6px; }
    ::-webkit-scrollbar-thumb { background: #C4C7C5; border-radius: 10px; }
    </style>
    """,
    unsafe_allow_html=True,
)

# =====================================================================
# ВСПАМОГАТЕЛЬНЫЕ ФУНКЦИИ БЕЗОПАСНОГО ПРИВЕДЕНИЯ ТИПОВ
# =====================================================================
def safe_float(val) -> float:
    if isinstance(val, list):
        val = val[0] if len(val) > 0 else 0
    try:
        return float(val) if val is not None else 0.0
    except (ValueError, TypeError):
        return 0.0

def safe_int(val) -> int:
    if isinstance(val, list):
        val = val[0] if len(val) > 0 else 0
    try:
        return int(val) if val is not None else 0
    except (ValueError, TypeError):
        return 0

def safe_str(val) -> str:
    if isinstance(val, list):
        return " ".join([str(x) for x in val])
    return str(val) if val is not None else ""

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

# =====================================================================
# 2. ФУНКЦИИ ГЕНЕРАЦИИ PDF-ОТЧЕТОВ
# =====================================================================
def generate_pdf_report(project_name: str, messages: list) -> bytes:
    font_path = "DejaVuSans.ttf"
    if not os.path.exists(font_path):
        try:
            urllib.request.urlretrieve(
                "https://cdn.jsdelivr.net/font-dejavu/2.37/ttf/DejaVuSans.ttf",
                font_path,
            )
        except Exception:
            pass

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=30,
        leftMargin=30,
        topMargin=30,
        bottomMargin=30,
    )
    font_name = "DejaVu" if os.path.exists(font_path) else "Helvetica"
    if font_name == "DejaVu":
        pdfmetrics.registerFont(TTFont("DejaVu", font_path))

    style_title = ParagraphStyle(
        "DocTitle", fontName=font_name, fontSize=16, leading=20, spaceAfter=12
    )
    style_meta = ParagraphStyle(
        "DocMeta",
        fontName=font_name,
        fontSize=9,
        leading=12,
        textColor="#555555",
        spaceAfter=18,
    )
    style_msg = ParagraphStyle(
        "DocMsg", fontName=font_name, fontSize=10, leading=14, spaceAfter=10
    )

    story = [
        Paragraph(f"<b>Отчет по диалогу: {project_name}</b>", style_title),
        Paragraph(
            f"Дата генерации: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |"
            " Сгенерировано системой Enterprise AI",
            style_meta,
        ),
    ]

    for msg in messages:
        role_label = (
            "👤 <b>Пользователь</b>"
            if msg["role"] == "user"
            else "🤖 <b>AI Ассистент</b>"
        )
        clean_content = (
            msg["content"]
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )
        story.append(Paragraph(f"{role_label}:<br/>{clean_content}", style_msg))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# =====================================================================
# 3. ФУНКЦИИ ИЗВЛЕЧЕНИЯ ТЕКСТА И ВЕКТОРИЗАЦИИ
# =====================================================================
def extract_text_from_file(uploaded_file) -> str:
    fname = uploaded_file.name.lower()
    try:
        if fname.endswith(".pdf"):
            pdf_reader = PdfReader(uploaded_file)
            return "\n\n".join(
                [page.extract_text() or "" for page in pdf_reader.pages]
            )
        elif fname.endswith(".docx"):
            doc = Document(uploaded_file)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = [
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    ]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            return "\n\n".join(paragraphs)
        else:
            return uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception as e:
        st.error(f"Ошибка при чтении файла {uploaded_file.name}: {e}")
        return ""


def split_text_into_chunks(text: str, chunk_size: int = 600) -> list:
    paragraphs = text.split("\n\n")
    chunks, current_chunk = [], ""
    for p in paragraphs:
        p_clean = p.strip()
        if not p_clean:
            continue
        if len(current_chunk) + len(p_clean) < chunk_size:
            current_chunk += p_clean + "\n\n"
        else:
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            current_chunk = p_clean + "\n\n"
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    return chunks if chunks else [text]


def get_cloud_embedding(text: str) -> list:
    if HF_TOKEN:
        try:
            client = InferenceClient(token=HF_TOKEN)
            result = client.feature_extraction(
                text,
                model="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
            )
            data = result.tolist() if hasattr(result, "tolist") else result
            while (
                isinstance(data, list) and len(data) > 0 and isinstance(data[0], list)
            ):
                data = data[0]
            return [float(x) for x in data]
        except Exception:
            pass
    try:
        from fastembed import TextEmbedding

        embed_model = TextEmbedding(
            model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
            threads=1,
        )
        return list(embed_model.embed([text]))[0].tolist()
    except Exception:
        return [0.0] * 384


# =====================================================================
# 4. GeoIP ФУНКЦИИ
# =====================================================================
def get_client_ip() -> str:
    try:
        if hasattr(st, "context") and hasattr(st.context, "headers"):
            headers = st.context.headers
            if "X-Forwarded-For" in headers:
                return headers["X-Forwarded-For"].split(",")[0].strip()
            if "X-Real-Ip" in headers:
                return headers["X-Real-Ip"]
    except Exception:
        pass
    return "127.0.0.1"


def get_geoip_details(ip: str) -> dict:
    default_res = {
        "country": "Неизвестно",
        "city": "Неизвестно",
        "lat": None,
        "lon": None,
    }
    if (
        ip in ["127.0.0.1", "localhost"]
        or ip.startswith("192.168.")
        or ip.startswith("10.")
    ):
        return {
            "country": "Локальная сеть",
            "city": "Dev",
            "lat": 50.4501,
            "lon": 30.5234,
        }
    try:
        url = f"http://ip-api.com/json/{ip}?fields=country,city,lat,lon,status"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=1.5) as response:
            data = json.loads(response.read().decode())
            if data.get("status") == "success":
                return {
                    "country": data.get("country", "Неизвестно"),
                    "city": data.get("city", "Неизвестно"),
                    "lat": data.get("lat"),
                    "lon": data.get("lon"),
                }
    except Exception:
        pass
    return default_res


# =====================================================================
# 5. ИНИЦИАЛИЗАЦИЯ СЕРВИСОВ
# =====================================================================
@st.cache_resource(max_entries=1)
def init_services():
    qdrant = QdrantClient(
        url=QDRANT_URL,
        api_key=QDRANT_API_KEY,
        port=443,
        https=True,
        check_compatibility=False,
    )
    collections = [c.name for c in qdrant.get_collections().collections]

    for c_name in [
        COLLECTION_NAME,
        LOGS_COLLECTION,
        ANALYTICS_COLLECTION,
        CONFIG_COLLECTION,
        CHAT_HISTORY_COLLECTION,
        RE_COLLECTION_NAME,
    ]:
        if c_name not in collections:
            size = (
                1
                if c_name in [ANALYTICS_COLLECTION, CHAT_HISTORY_COLLECTION]
                else 384
            )
            qdrant.create_collection(
                collection_name=c_name,
                vectors_config=VectorParams(size=size, distance=Distance.COSINE),
            )

    for col in [COLLECTION_NAME, CHAT_HISTORY_COLLECTION, RE_COLLECTION_NAME]:
        for field in [
            "section",
            "project",
            "source_file",
            "username",
            "deal_type",
            "district",
            "post_type",
        ]:
            try:
                qdrant.create_payload_index(
                    collection_name=col,
                    field_name=field,
                    field_schema=PayloadSchemaType.KEYWORD,
                )
            except Exception:
                pass

    groq_client = Groq(api_key=GROQ_API_KEY)
    return qdrant, groq_client


qdrant, groq_client = init_services()


# =====================================================================
# 6. ФУНКЦИИ УПРАВЛЕНИЯ ЧАТАМИ И КОНФИГУРАЦИЕЙ
# =====================================================================
def load_system_config():
    try:
        scroll_res, _ = qdrant.scroll(
            collection_name=CONFIG_COLLECTION, limit=5, with_payload=True
        )
        for pt in scroll_res:
            if pt.payload and "projects" in pt.payload:
                return pt.payload.get("projects"), pt.payload.get("sections")
    except Exception:
        pass
    return None, None


def save_system_config(projects, sections):
    try:
        point = PointStruct(
            id="00000000-0000-0000-0000-000000000001",
            vector=[0.0] * 384,
            payload={"projects": projects, "sections": list(set(sections))},
        )
        qdrant.upsert(collection_name=CONFIG_COLLECTION, points=[point])
    except Exception as e:
        print(f"Ошибка сохранения конфига: {e}")


def get_recent_chat_threads(username: str, limit: int = 50) -> list:
    try:
        scroll_res, _ = qdrant.scroll(
            collection_name=CHAT_HISTORY_COLLECTION,
            limit=limit,
            with_payload=True,
            with_vectors=False,
        )
        threads = []
        for pt in scroll_res:
            p = pt.payload or {}
            if p.get("username") == username:
                threads.append({
                    "chat_id": str(pt.id),
                    "title": p.get("title", "Новый чат"),
                    "project": p.get("project", "Общий проект"),
                    "updated_at": p.get("updated_at", ""),
                    "messages": p.get("messages", []),
                    "is_pinned": p.get("is_pinned", False),
                })
        return sorted(
            [t for t in threads if t.get("is_pinned")],
            key=lambda x: x.get("updated_at", ""),
            reverse=True,
        ) + sorted(
            [t for t in threads if not t.get("is_pinned")],
            key=lambda x: x.get("updated_at", ""),
            reverse=True,
        )
    except Exception:
        return []


def load_chat_thread_by_id(chat_id: str) -> tuple:
    try:
        valid_uuid = (
            str(uuid.UUID(chat_id))
            if len(chat_id) == 36
            else str(uuid.uuid5(uuid.NAMESPACE_DNS, chat_id))
        )
        res = qdrant.retrieve(
            collection_name=CHAT_HISTORY_COLLECTION,
            ids=[valid_uuid],
            with_payload=True,
        )
        if res and res[0].payload:
            p = res[0].payload
            return (
                p.get("messages", []),
                p.get("project", "Общий проект"),
                p.get("title", "Диалог"),
            )
    except Exception:
        pass
    return [], "Общий проект", "Новый чат"


def save_chat_thread(
    chat_id: str, username: str, project: str, title: str, messages: list
):
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    valid_uuid = (
        str(uuid.UUID(chat_id))
        if len(chat_id) == 36
        else str(uuid.uuid5(uuid.NAMESPACE_DNS, chat_id))
    )
    payload = {
        "chat_id": chat_id,
        "username": username,
        "project": project,
        "title": title,
        "messages": messages,
        "updated_at": now_str,
    }
    try:
        qdrant.upsert(
            collection_name=CHAT_HISTORY_COLLECTION,
            points=[PointStruct(id=valid_uuid, vector=[0.0], payload=payload)],
        )
    except Exception:
        pass


def toggle_pin_chat_thread(chat_id: str, is_pinned: bool):
    try:
        valid_uuid = (
            str(uuid.UUID(chat_id))
            if len(chat_id) == 36
            else str(uuid.uuid5(uuid.NAMESPACE_DNS, chat_id))
        )
        qdrant.set_payload(
            collection_name=CHAT_HISTORY_COLLECTION,
            payload={"is_pinned": is_pinned},
            points=[valid_uuid],
        )
    except Exception:
        pass


def rename_chat_thread(chat_id: str, new_title: str):
    try:
        valid_uuid = (
            str(uuid.UUID(chat_id))
            if len(chat_id) == 36
            else str(uuid.uuid5(uuid.NAMESPACE_DNS, chat_id))
        )
        qdrant.set_payload(
            collection_name=CHAT_HISTORY_COLLECTION,
            payload={"title": new_title},
            points=[valid_uuid],
        )
        if st.session_state.active_chat_id == chat_id:
            st.session_state.active_chat_title = new_title
    except Exception:
        pass


def delete_chat_thread(chat_id: str):
    try:
        valid_uuid = (
            str(uuid.UUID(chat_id))
            if len(chat_id) == 36
            else str(uuid.uuid5(uuid.NAMESPACE_DNS, chat_id))
        )
        qdrant.delete(
            collection_name=CHAT_HISTORY_COLLECTION, points_selector=[valid_uuid]
        )
    except Exception:
        pass


def log_event(
    action: str,
    details: str,
    ip: str = None,
    username: str = None,
    role: str = None,
):
    try:
        user_info = st.session_state.get("current_user") or {}
        req_username = username if username else user_info.get("username", "Гость")
        req_role = role if role else user_info.get("role", "guest")
        req_ip = ip if ip else user_info.get("ip", get_client_ip())
        geo = get_geoip_details(req_ip)
        log_point = PointStruct(
            id=uuid.uuid4().hex,
            vector=[0.0] * 384,
            payload={
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "username": req_username,
                "role": req_role,
                "action": action,
                "details": details,
                "ip": req_ip,
                "country": geo.get("country"),
                "city": geo.get("city"),
                "lat": geo.get("lat"),
                "lon": geo.get("lon"),
            },
        )
        qdrant.upsert(collection_name=LOGS_COLLECTION, points=[log_point])
    except Exception:
        pass


def log_analytics(
    source: str,
    user_id: str,
    username: str,
    event_type: str,
    query: str = "",
    score: float = 0.0,
    status: str = "Успешно",
    details: str = "",
):
    try:
        point = PointStruct(
            id=uuid.uuid4().hex,
            vector=[0.0],
            payload={
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "source": source,
                "user_id": str(user_id),
                "username": username or "Аноним",
                "event_type": event_type,
                "query": query[:300],
                "score": float(score),
                "found_in_kb": bool(score >= 0.20),
                "status": status,
                "details": details,
            },
        )
        qdrant.upsert(collection_name=ANALYTICS_COLLECTION, points=[point])
    except Exception:
        pass


def get_audit_logs():
    try:
        scroll_res, _ = qdrant.scroll(
            collection_name=LOGS_COLLECTION,
            limit=1000,
            with_payload=True,
            with_vectors=False,
        )
        return sorted(
            [pt.payload for pt in scroll_res if pt.payload],
            key=lambda x: x.get("timestamp", ""),
            reverse=True,
        )
    except Exception:
        return []


def get_db_files_summary():
    try:
        scroll_res, _ = qdrant.scroll(
            collection_name=COLLECTION_NAME,
            limit=10000,
            with_payload=["source_file", "section"],
            with_vectors=False,
        )
        files_by_section = {}
        for point in scroll_res:
            sec = point.payload.get("section", "Общий раздел")
            src = point.payload.get("source_file", "Неизвестный файл")
            if sec not in files_by_section:
                files_by_section[sec] = {}
            files_by_section[sec][src] = files_by_section[sec].get(src, 0) + 1
        return files_by_section
    except Exception:
        return {}


# =====================================================================
# РАСШИРЕННЫЙ ФИЛЬТР МУСОРА И ОБЯЗАТЕЛЬНАЯ ПРОВЕРКА ТЕМЫ НЕДВИЖИМОСТИ
# =====================================================================
NON_REAL_ESTATE_KEYWORDS = [
    "запрещенное слово", "ограничен(а)", "выключил звук", "підробіток", "подработка",
    "приєднуйся до команди", "дохід", "доход", "оплата від", "оплата от", "грн на день",
    "грн в день", "грн на місяць", "грн в месяц", "вакансія", "вакансия", "требуются",
    "потрібні", "работу", "робота", "зарплата", "заработок", "заробіток", "клининг",
    "клинінг", "консультаци", "консультаці", "петици", "петиці", "инстаграм", "instagram",
    "репетитор", "майстер", "майстра", "сантехн", "электрик", "електрик", "няня", "няню",
    "массаж", "масаж", "маникюр", "манікюр", "стрижк", "авто", "машин", "перевозк",
    "перевезення", "психолог", "лікар", "врач", "дитини", "ребенка", "школьн", "урок",
    "занятия", "заняття", "дефіцит", "можливо й так",
]

REAL_ESTATE_KEYWORDS = [
    "квартир", "будинок", "дом", "кімнат", "комнат", "приміщен", "помещен", "офіс",
    "офис", "ділянк", "участок", "жиль", "житл", "жк", "снять", "знять", "винайм",
    "аренд", "оренд", "продам", "продаж", "сдам", "здам", "койко",
]


def determine_post_type_dynamically(payload: dict) -> str:
    raw_text = safe_str(payload.get("raw_text")).lower()
    parsed = payload.get("parsed_data") or {}

    # 1. Отсеиваем черный список бытовых услуг и диалогов
    for tk in NON_REAL_ESTATE_KEYWORDS:
        if tk in raw_text:
            return "trash"

    # 2. Проверяем обязательное наличие ключевых слов недвижимости
    has_re_terms = any(rk in raw_text for rk in REAL_ESTATE_KEYWORDS)
    has_parsed_specs = (
        safe_float(parsed.get("price_usd")) > 0
        or safe_int(parsed.get("rooms")) > 0
        or safe_float(parsed.get("area_sqm")) > 0
    )

    if not has_re_terms and not has_parsed_specs:
        return "trash"

    # 3. Разделяем предложения и запросы клиентов
    is_demand = any(
        dk in raw_text
        for dk in [
            "шукаю", "ищу", "потрібно", "нужно", "підшукуємо", "подбираем",
            "для своїх клієнтів", "для своих клиентов", "винайму", "сниму",
            "снять", "купим", "купити", "купимо", "купить",
        ]
    )
    if "жильцов" in raw_text or "орендарів" in raw_text:
        is_demand = False

    if is_demand:
        return "demand"

    stored_type = payload.get("post_type") or parsed.get("post_type")
    if stored_type in ["demand", "trash"]:
        return stored_type

    return "offer"


def fetch_real_estate_listings(
    post_type="offer",
    deal_type="Все",
    property_type="Все",
    min_price=0,
    max_price=500000,
    min_area=0,
    max_area=1000,
    rooms_filter="Все",
    owner_only=False,
    district_query="",
):
    try:
        must_conditions = []
        if deal_type != "Все":
            deal_val = (
                "rent" if ("Снять" in deal_type or "Аренд" in deal_type) else "sale"
            )
            must_conditions.append(
                FieldCondition(key="deal_type", match=MatchValue(value=deal_val))
            )

        scroll_filter = Filter(must=must_conditions) if must_conditions else None

        points, _ = qdrant.scroll(
            collection_name=RE_COLLECTION_NAME,
            limit=600,
            scroll_filter=scroll_filter,
            with_payload=True,
            with_vectors=False,
        )

        results = []

        for pt in points:
            p = pt.payload or {}
            parsed = p.get("parsed_data") or {}

            item_post_type = determine_post_type_dynamically(p)

            if item_post_type != post_type or item_post_type == "trash":
                continue

            price = safe_float(parsed.get("price_usd"))
            area = safe_float(parsed.get("area_sqm"))
            rooms = parsed.get("rooms")

            district = safe_str(parsed.get("district")).lower()
            address = safe_str(parsed.get("address")).lower()
            raw_text = safe_str(p.get("raw_text")).lower()
            is_broker = parsed.get("is_broker")

            if owner_only and is_broker is not False:
                continue
            if price > 0 and (price < min_price or price > max_price):
                continue
            if area > 0 and (area < min_area or area > max_area):
                continue

            if property_type != "Все":
                prop_key = property_type.lower()[:4]
                combined_text = f"{district} {address} {raw_text}"
                if prop_key not in combined_text:
                    continue

            if rooms_filter != "Все":
                rooms_cnt = safe_int(rooms)
                if rooms_filter == "4+":
                    if rooms_cnt < 4:
                        continue
                else:
                    try:
                        if rooms_cnt != int(rooms_filter):
                            continue
                    except ValueError:
                        pass

            if district_query.strip():
                dq = district_query.strip().lower()
                if dq not in district and dq not in address and dq not in raw_text:
                    continue

            results.append(p)

        return sorted(
            results, key=lambda x: safe_str(x.get("created_at")), reverse=True
        )
    except Exception as e:
        st.error(f"⚠️ Ошибка при выгрузке объектов: {e}")
        return []


# =====================================================================
# СЕССИИ И БД ПОЛЬЗОВАТЕЛЕЙ
# =====================================================================
if "users_db" not in st.session_state:
    st.session_state.users_db = {
        "owner": {
            "password": hash_password("owner123"),
            "role": "owner",
            "name": "Олексій Марфенков",
            "failed_attempts": 0,
            "is_blocked": False,
            "max_connections": 5,
            "active_sessions": 0,
        },
        "admin": {
            "password": hash_password("admin123"),
            "role": "admin",
            "name": "Администратор",
            "failed_attempts": 0,
            "is_blocked": False,
            "max_connections": 3,
            "active_sessions": 0,
        },
        "user": {
            "password": hash_password("user123"),
            "role": "user",
            "name": "Менеджер",
            "failed_attempts": 0,
            "is_blocked": False,
            "max_connections": 1,
            "active_sessions": 0,
        },
    }

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "view_mode" not in st.session_state:
    st.session_state.view_mode = "chat"

if "projects" not in st.session_state or "sections" not in st.session_state:
    db_projects, db_sections = load_system_config()
    if db_projects and db_sections:
        st.session_state.projects = db_projects
        st.session_state.sections = db_sections
    else:
        st.session_state.sections = [
            "Общий раздел",
            "Продажи и CRM",
            "Регламенты",
            "Техническая часть",
        ]
        st.session_state.projects = {
            "Общий проект": ["Общий раздел"],
            "Creatio 2.0": ["Продажи и CRM", "Общий раздел"],
            "КиберБез": ["Техническая часть", "Регламенты"],
        }
        save_system_config(st.session_state.projects, st.session_state.sections)

if "active_chat_id" not in st.session_state:
    st.session_state.active_chat_id = str(uuid.uuid4())
if "active_chat_title" not in st.session_state:
    st.session_state.active_chat_title = "Новый чат"
if "messages" not in st.session_state:
    st.session_state.messages = [{
        "role": "assistant",
        "content": "Здравствуйте! Чем я могу помочь вам сегодня?",
    }]
if "metrics_history" not in st.session_state:
    st.session_state.metrics_history = []
if "voice_key_counter" not in st.session_state:
    st.session_state.voice_key_counter = 0

if st.session_state.logged_in:
    current_time = time.time()
    last_act = st.session_state.get("last_activity_time", current_time)
    if (current_time - last_act) > (SESSION_TIMEOUT_MINUTES * 60):
        st.session_state.logged_in = False
        st.session_state.current_user = None
        st.session_state.timeout_message = (
            "⏳ Сессия завершена автоматически из-за неактивности"
            f" >{SESSION_TIMEOUT_MINUTES} мин."
        )
        st.rerun()
    else:
        st.session_state.last_activity_time = current_time

if not st.session_state.logged_in:
    col_l1, col_l2, col_l3 = st.columns([1, 2, 1])
    with col_l2:
        st.markdown(
            "<h1 style='text-align: center;'>🤖 Вход в Mavbot AI</h1>",
            unsafe_allow_html=True,
        )
        if st.session_state.get("timeout_message"):
            st.warning(st.session_state["timeout_message"])
        client_ip = get_client_ip()
        geo_info = get_geoip_details(client_ip)
        st.info(
            f"🌐 Ваш IP: `{client_ip}` | Страна:"
            f" **{geo_info['country']}** ({geo_info['city']})"
        )

        with st.form("login_form"):
            user_input = st.text_input("Логин:")
            pass_input = st.text_input("Пароль:", type="password")
            if st.form_submit_button("Войти в систему", use_container_width=True):
                clean_user = user_input.strip().lower()
                user_record = st.session_state.users_db.get(clean_user)
                if not user_record or user_record["password"] != hash_password(
                    pass_input
                ):
                    st.error("Неверный логин или пароль")
                else:
                    st.session_state.logged_in = True
                    st.session_state.last_activity_time = time.time()
                    st.session_state.current_user = {
                        "username": clean_user,
                        "role": user_record["role"],
                        "name": user_record["name"],
                        "ip": client_ip,
                    }
                    st.rerun()
    st.stop()

# =====================================================================
# SIDEBAR
# =====================================================================
user_data = st.session_state.current_user
user_role = user_data["role"]
role_badges = {
    "owner": "👑 Собственник",
    "admin": "🛠️ Администратор",
    "user": "👤 Пользователь",
}

with st.sidebar:
    st.markdown(
        '<div class="mavbot-header"><span class="mavbot-icon">🤖</span>'
        " <span>Mavbot</span></div>",
        unsafe_allow_html=True,
    )
    col_nav1, col_nav2 = st.columns([1, 1])
    with col_nav1:
        if st.button("➕ Чат", use_container_width=True, key="btn_new_chat"):
            st.session_state.view_mode = "chat"
            st.session_state.active_chat_id = str(uuid.uuid4())
            st.session_state.active_chat_title = "Новый чат"
            st.session_state.messages = [{
                "role": "assistant",
                "content": "Здравствуйте! Чем я могу помочь вам сегодня?",
            }]
            st.rerun()
    with col_nav2:
        btn_re_label = (
            "🏠 Недвижимость"
            if st.session_state.view_mode != "real_estate"
            else "📌 🏠 Недвижимость"
        )
        if st.button(
            btn_re_label, use_container_width=True, key="btn_real_estate_mode"
        ):
            st.session_state.view_mode = "real_estate"
            st.rerun()

    selected_project = st.selectbox(
        "Выберите проект:",
        list(st.session_state.projects.keys()),
        label_visibility="collapsed",
    )
    st.session_state.selected_project = selected_project

    st.markdown(
        '<div class="recent-title">Недавние</div>', unsafe_allow_html=True
    )
    recent_threads = get_recent_chat_threads(user_data["username"])
    for thread in recent_threads[:15]:
        t_id, t_title, is_pinned = (
            thread["chat_id"],
            thread["title"],
            thread.get("is_pinned", False),
        )
        prefix = "📌 " if is_pinned else "💬 "
        if st.button(
            f"{prefix}{t_title[:22]}", key=f"rec_{t_id}", use_container_width=True
        ):
            st.session_state.view_mode = "chat"
            st.session_state.active_chat_id = t_id
            msgs, _, title_loaded = load_chat_thread_by_id(t_id)
            st.session_state.messages = (
                msgs if msgs else [{"role": "assistant", "content": "Здравствуйте!"}]
            )
            st.session_state.active_chat_title = title_loaded
            st.rerun()

    st.markdown("---")
    if user_role in ["admin", "owner"]:
        if st.button(
            "⚙️ Настройки", use_container_width=True, key="btn_settings_sidebar"
        ):
            st.session_state.view_mode = "settings"
            st.rerun()

    st.markdown(
        f'<div class="user-profile-card"><div'
        f' class="user-avatar">{user_data["name"][0].upper()}</div><div><b>{user_data["name"]}</b><br/><small>{role_badges.get(user_role, user_role)}</small></div></div>',
        unsafe_allow_html=True,
    )
    if st.button("🚪 Выйти", use_container_width=True):
        st.session_state.logged_in = False
        st.rerun()

# =====================================================================
# MAIN WORKSPACE
# =====================================================================
if st.session_state.view_mode == "chat":
    st.title(f"🤖 Mavbot — [{selected_project}]")
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])

    text_prompt = st.chat_input(
        f"Спросите что-нибудь по проекту '{selected_project}'..."
    )
    if text_prompt:
        st.session_state.messages.append({"role": "user", "content": text_prompt})
        query_vector = get_cloud_embedding(text_prompt)
        res = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": text_prompt}],
            temperature=0.1,
        )
        answer = res.choices[0].message.content
        st.session_state.messages.append({"role": "assistant", "content": answer})
        save_chat_thread(
            st.session_state.active_chat_id,
            user_data["username"],
            selected_project,
            st.session_state.active_chat_title,
            st.session_state.messages,
        )
        st.rerun()

elif st.session_state.view_mode == "real_estate":
    st.title("🏠 Мониторинг Недвижимости Telegram")
    st.caption(
        "Автоматическая жесткая фильтрация бытовых услуг и разделение предложений"
        " объектов и клиентских запросов."
    )

    re_tabs = st.tabs([
        "🏢 Предложения (Объекты)",
        "🔍 Запросы клиентов (Ищут жилье)",
        "🤖 AI-Подбор под запрос",
    ])

    for tab_idx, target_type in enumerate(["offer", "demand"]):
        with re_tabs[tab_idx]:
            col_f1, col_f2, col_f3, col_f4, col_f5 = st.columns(5)
            with col_f1:
                with st.popover(
                    "🏡 Источник / Сделка ˅",
                    use_container_width=True,
                    key=f"p1_{target_type}",
                ):
                    f_owner_only = st.checkbox(
                        "🏡 Только от хозяина", value=False, key=f"re_owner_{target_type}"
                    )
                    f_deal_type = st.radio(
                        "Тип сделки:",
                        ["Все", "Снять", "Купить"],
                        key=f"re_deal_{target_type}",
                    )

            with col_f2:
                with st.popover(
                    "🏠 Тип объекта ˅",
                    use_container_width=True,
                    key=f"p2_{target_type}",
                ):
                    f_prop_type = st.selectbox(
                        "Тип:",
                        [
                            "Все",
                            "Квартира",
                            "Дом",
                            "Комната",
                            "Коммерческая",
                            "Участок",
                        ],
                        key=f"re_prop_{target_type}",
                    )

            with col_f3:
                with st.popover(
                    "💰 Цена ($) ˅", use_container_width=True, key=f"p3_{target_type}"
                ):
                    f_min_price = st.number_input(
                        "От ($):",
                        min_value=0,
                        max_value=500000,
                        value=0,
                        key=f"re_pmin_{target_type}",
                    )
                    f_max_price = st.number_input(
                        "До ($):",
                        min_value=0,
                        max_value=1000000,
                        value=500000,
                        key=f"re_pmax_{target_type}",
                    )

            with col_f4:
                with st.popover(
                    "🚪 Комнаты ˅", use_container_width=True, key=f"p4_{target_type}"
                ):
                    f_rooms = st.radio(
                        "Комнат:",
                        ["Все", "1", "2", "3", "4+"],
                        key=f"re_rooms_{target_type}",
                    )

            with col_f5:
                with st.popover(
                    "⚙️ Фильтры ˅", use_container_width=True, key=f"p5_{target_type}"
                ):
                    f_district = st.text_input(
                        "Район / Улица:", key=f"re_dist_{target_type}"
                    )

            listings = fetch_real_estate_listings(
                post_type=target_type,
                deal_type=f_deal_type,
                property_type=f_prop_type,
                min_price=f_min_price,
                max_price=f_max_price,
                rooms_filter=f_rooms,
                owner_only=f_owner_only,
                district_query=f_district,
            )

            st.markdown(f"Найдено чистых целевых записей: **{len(listings)}**")
            st.divider()

            if not listings:
                st.info("Записи не найдены.")
            else:
                for item in listings:
                    parsed = item.get("parsed_data") or {}
                    price_val = safe_float(parsed.get("price_usd"))
                    price_str = (
                        f"${int(price_val):,}" if price_val > 0 else "Цена не указана"
                    )

                    deal_lbl = (
                        "Аренда" if parsed.get("deal_type") == "rent" else "Продажа"
                    )
                    rooms_val = safe_int(parsed.get("rooms"))
                    rooms_lbl = f"{rooms_val} к." if rooms_val > 0 else "Комнаты не указаны"
                    district_lbl = parsed.get("district") or "Район не указан"
                    address_lbl = parsed.get("address") or ""
                    phone_lbl = parsed.get("phone") or "Телефон не указан"
                    is_broker = parsed.get("is_broker")

                    broker_tag = (
                        '<span class="re-badge-owner">🏡 От хозяина</span>'
                        if is_broker is False
                        else '<span class="re-badge">👔 Риелтор / Агентство</span>'
                    )

                    card_html = f"""<div class="re-card">
<div style="display: flex; justify-content: space-between; align-items: center;">
<span class="re-price">{price_str}</span>
<div>
<span class="re-badge">{deal_lbl}</span>
<span class="re-badge">{rooms_lbl}</span>
{broker_tag}
</div>
</div>
<div style="margin-top: 8px; font-weight: 600; color: #3c4043;">📍 {district_lbl} {f'— {address_lbl}' if address_lbl else ''}</div>
<div style="margin-top: 4px; font-size: 13px; color: #5f6368;">📞 {phone_lbl} | 💬 Канал: <b>{item.get('channel', 'Telegram')}</b> | 🕒 {item.get('created_at', '')}</div>
</div>"""
                    st.markdown(card_html, unsafe_allow_html=True)
                    with st.expander("📄 Описание из Telegram"):
                        st.text(item.get("raw_text", ""))

    with re_tabs[2]:
        st.subheader("🤖 AI Поиск объектов")
        ai_re_prompt = st.text_input("Поисковый запрос:")
        if ai_re_prompt and st.button("Искать", use_container_width=True):
            st.success("Интеллектуальный поиск активен.")

else:
    col_head1, col_head2 = st.columns([4, 1])
    with col_head1:
        st.title(f"⚙️ Настройки системы — [{selected_project}]")
    with col_head2:
        if st.button("💬 Вернуться в чат", use_container_width=True):
            st.session_state.view_mode = "chat"
            st.rerun()

    settings_tab_titles = []
    if user_role in ["admin", "owner"]:
        settings_tab_titles.extend(
            ["📁 Загрузка документов", "🗂️ Управление файлами", "📈 Аналитика"]
        )
    if user_role == "owner":
        settings_tab_titles.extend([
            "📜 Полный Журнал Логов",
            "🗺️ Карта Входов (GeoIP)",
            "💡 Пробелы в знаниях & Отзывы",
            "👥 Управление Аккаунтами",
        ])

    s_tabs = st.tabs(settings_tab_titles)
    s_tab_dict = {title: tab for title, tab in zip(settings_tab_titles, s_tabs)}

    if "📁 Загрузка документов" in s_tab_dict:
        with s_tab_dict["📁 Загрузка документов"]:
            st.subheader("📁 Пополнение Базы Знаний (PDF, Word, Text, Markdown)")
            col_up1, col_up2 = st.columns([2, 1])

            with col_up1:
                target_section = st.selectbox(
                    "Целевой раздел:", st.session_state.sections
                )
            with col_up2:
                new_sec_input = st.text_input("➕ Новый раздел:")
                if st.button("Добавить раздел", use_container_width=True):
                    if new_sec_input and new_sec_input not in st.session_state.sections:
                        st.session_state.sections.append(new_sec_input)
                        save_system_config(
                            st.session_state.projects, st.session_state.sections
                        )
                        log_event("CREATE_SECTION", f"Создан раздел '{new_sec_input}'")
                        st.success(f"Раздел '{new_sec_input}' создан!")
                        st.rerun()

            st.divider()
            uploaded_files = st.file_uploader(
                "Перетащите файлы (`.pdf`, `.docx`, `.txt`, `.md`):",
                type=["pdf", "docx", "txt", "md"],
                accept_multiple_files=True,
            )

            if uploaded_files and st.button(
                f"🚀 Векторизовать и загрузить в '{target_section}'",
                use_container_width=True,
            ):
                markdown_splitter = MarkdownHeaderTextSplitter(
                    headers_to_split_on=[
                        ("#", "Header 1"),
                        ("##", "Header 2"),
                        ("###", "Header 3"),
                    ],
                    strip_headers=False,
                )

                all_points = []
                with st.spinner(
                    "Извлечение текста, нарезка на чанки и векторизация..."
                ):
                    for file in uploaded_files:
                        fname = file.name
                        extracted_text = extract_text_from_file(file)

                        if not extracted_text.strip():
                            st.warning(
                                f"Файл '{fname}' пуст или из него не удалось извлечь текст."
                            )
                            continue

                        if fname.lower().endswith(".md"):
                            chunks_md = markdown_splitter.split_text(extracted_text)
                            texts = (
                                [c.page_content for c in chunks_md]
                                if chunks_md
                                else [extracted_text]
                            )
                            metadatas = (
                                [c.metadata for c in chunks_md] if chunks_md else [{}]
                            )
                        else:
                            texts = split_text_into_chunks(extracted_text)
                            metadatas = [{}] * len(texts)

                        for idx, text_chunk in enumerate(texts):
                            emb = get_cloud_embedding(text_chunk)
                            all_points.append(
                                PointStruct(
                                    id=uuid.uuid4().hex,
                                    vector=emb,
                                    payload={
                                        "text": text_chunk,
                                        "source_file": fname,
                                        "section": target_section,
                                        **metadatas[idx],
                                    },
                                )
                            )

                    if all_points:
                        qdrant.upsert(collection_name=COLLECTION_NAME, points=all_points)
                        log_event(
                            "UPLOAD_FILES",
                            f"Загружено {len(uploaded_files)} файлов ({len(all_points)}"
                            f" чанков) в раздел '{target_section}'",
                        )
                        log_analytics(
                            "Web",
                            user_data.get("username"),
                            user_data.get("name"),
                            "Загрузка документа",
                            f"Файлов: {len(uploaded_files)}",
                            score=1.0,
                            status="Загружено",
                        )

                        st.success(
                            f"🎉 Успешно векторизовано файлов: {len(uploaded_files)} (всего"
                            f" {len(all_points)} чанков)!"
                        )
                        st.rerun()

    if "🗂️ Управление файлами" in s_tab_dict:
        with s_tab_dict["🗂️ Управление файлами"]:
            st.subheader("🗂️ Управление документами")
            files_by_sec = get_db_files_summary()

            if not files_by_sec:
                st.info("Файлы отсутствуют.")
            else:
                for sec_name, files_dict in files_by_sec.items():
                    with st.expander(
                        f"📁 Раздел: **{sec_name}** ({len(files_dict)} файлов)",
                        expanded=True,
                    ):
                        for fname, chunk_cnt in files_dict.items():
                            c1, c2 = st.columns([3, 1])
                            with c1:
                                st.write(f"📄 **{fname}** (`{chunk_cnt} чанков`)")
                                other_secs = [
                                    s for s in st.session_state.sections if s != sec_name
                                ]
                                if other_secs:
                                    dest_s = st.selectbox(
                                        "Переместить в:", other_secs, key=f"s_{sec_name}_{fname}"
                                    )
                                    if st.button("🚚 Переместить", key=f"m_{sec_name}_{fname}"):
                                        pts, _ = qdrant.scroll(
                                            collection_name=COLLECTION_NAME,
                                            scroll_filter=Filter(
                                                must=[
                                                    FieldCondition(
                                                        key="source_file",
                                                        match=MatchValue(value=fname),
                                                    ),
                                                    FieldCondition(
                                                        key="section",
                                                        match=MatchValue(value=sec_name),
                                                    ),
                                                ]
                                            ),
                                            limit=10000,
                                            with_payload=False,
                                            with_vectors=False,
                                        )
                                        p_ids = [p.id for p in pts]
                                        if p_ids:
                                            qdrant.set_payload(
                                                collection_name=COLLECTION_NAME,
                                                payload={"section": dest_s},
                                                points=p_ids,
                                            )
                                            log_event(
                                                "MOVE_FILE",
                                                f"Файл '{fname}' из '{sec_name}' в '{dest_s}'",
                                            )
                                            st.success("Перемещено!")
                                            st.rerun()

                            with c2:
                                if st.button(
                                    "🗑️ Удалить", key=f"d_{sec_name}_{fname}", type="primary"
                                ):
                                    pts, _ = qdrant.scroll(
                                        collection_name=COLLECTION_NAME,
                                        scroll_filter=Filter(
                                            must=[
                                                FieldCondition(
                                                    key="source_file", match=MatchValue(value=fname)
                                                ),
                                                FieldCondition(
                                                    key="section", match=MatchValue(value=sec_name)
                                                ),
                                            ]
                                        ),
                                        limit=10000,
                                        with_payload=False,
                                        with_vectors=False,
                                    )
                                    p_ids = [p.id for p in pts]
                                    if p_ids:
                                        qdrant.delete(
                                            collection_name=COLLECTION_NAME,
                                            points_selector=p_ids,
                                        )
                                        log_event(
                                            "DELETE_FILE",
                                            f"Файл '{fname}' удален из '{sec_name}'",
                                        )
                                        st.success("Удалено!")
                                        st.rerun()
                            st.divider()

    if "📈 Аналитика" in s_tab_dict:
        with s_tab_dict["📈 Аналитика"]:
            st.subheader("📈 Статистика использования системы")

            c_ref, _ = st.columns([1, 4])
            with c_ref:
                if st.button("🔄 Обновить данные аналитики", use_container_width=True):
                    st.rerun()

            try:
                scroll_res, _ = qdrant.scroll(
                    collection_name=ANALYTICS_COLLECTION,
                    limit=1000,
                    with_payload=True,
                    with_vectors=False,
                )

                if scroll_res:
                    analytics_data = [pt.payload for pt in scroll_res if pt.payload]
                    df_a = pd.DataFrame(analytics_data)

                    total_q = len(df_a)
                    tg_q = (
                        len(df_a[df_a["source"] == "Telegram"])
                        if "source" in df_a.columns
                        else 0
                    )
                    web_q = (
                        len(df_a[df_a["source"] == "Web"])
                        if "source" in df_a.columns
                        else 0
                    )

                    success_count = (
                        len(df_a[df_a["found_in_kb"] == True])
                        if "found_in_kb" in df_a.columns
                        else 0
                    )
                    success_pct = (
                        round((success_count / total_q) * 100, 1) if total_q > 0 else 0
                    )

                    m1, m2, m3, m4 = st.columns(4)
                    m1.metric("Всего обращений", total_q)
                    m2.metric("Из Telegram 📱", tg_q)
                    m3.metric("Из Веб-чата 🌐", web_q)
                    m4.metric("Найдено в БЗ 🎯", f"{success_pct}%")

                    st.divider()
                    col_ch1, col_ch2 = st.columns(2)

                    with col_ch1:
                        st.markdown("### 📱 Распределение по источникам")
                        if "source" in df_a.columns:
                            st.bar_chart(df_a["source"].value_counts())

                    with col_ch2:
                        st.markdown("### 📊 Типы запросов и действий")
                        if "event_type" in df_a.columns:
                            st.bar_chart(df_a["event_type"].value_counts())

                    st.divider()
                    st.markdown("### 📜 Подробный журнал операций (Telegram & Web)")

                    if "timestamp" in df_a.columns:
                        df_a = df_a.sort_values(by="timestamp", ascending=False)

                    show_cols = [
                        c
                        for c in [
                            "timestamp",
                            "source",
                            "username",
                            "event_type",
                            "query",
                            "score",
                            "status",
                        ]
                        if c in df_a.columns
                    ]

                    st.dataframe(
                        df_a[show_cols],
                        column_config={
                            "timestamp": "Время (UTC)",
                            "source": "Источник",
                            "username": "Пользователь",
                            "event_type": "Тип действия",
                            "query": "Запрос / Файл",
                            "score": "Точность (Score)",
                            "status": "Результат",
                        },
                        use_container_width=True,
                        hide_index=True,
                    )
                else:
                    st.info("Пока нет зафиксированных данных в аналитике Qdrant.")

            except Exception as e:
                st.warning(f"Не удалось выгрузить данные из базы аналитики: {e}")

            if st.session_state.metrics_history:
                st.divider()
                st.markdown(
                    "### ⚡ Метрики скорости и токенов текущей веб-сессии (Groq + Qdrant)"
                )

                df_m = pd.DataFrame(st.session_state.metrics_history)

                total_reqs = len(df_m)
                total_tokens = df_m["Всего токенов"].sum()
                avg_time = round(df_m["Время ответа (сек)"].mean(), 2)
                avg_qdrant = round(df_m["Поиск Qdrant (мс)"].mean(), 0)

                GROQ_DAILY_LIMIT = 100000
                tokens_used_pct = round((total_tokens / GROQ_DAILY_LIMIT) * 100, 2)
                tokens_remaining = GROQ_DAILY_LIMIT - total_tokens

                col_m1, col_m2, col_m3, col_m4, col_m5 = st.columns(5)
                col_m1.metric("Запросов в сессии", total_reqs)
                col_m2.metric("Токенов за сессию", f"{total_tokens:,}")
                col_m3.metric(
                    "Остаток Groq TPD",
                    f"{tokens_remaining:,}",
                    f"-{tokens_used_pct}% лимита",
                    delta_color="normal",
                )
                col_m4.metric("Средний ответ LLM", f"{avg_time} с")
                col_m5.metric("Средний поиск Qdrant", f"{avg_qdrant:.0f} мс")

                st.caption(
                    f"📊 Расход суточного лимита Groq: **{total_tokens:,}** из **100,000**"
                    f" токенов ({tokens_used_pct}%):"
                )
                st.progress(min(total_tokens / GROQ_DAILY_LIMIT, 1.0))

                st.markdown("---")

                col_s1, col_s2 = st.columns(2)

                with col_s1:
                    st.markdown(
                        "##### 📊 Расход токенов по запросам (Prompt vs Generation)"
                    )
                    st.bar_chart(
                        df_m.set_index("Запрос №")[
                            ["Входные токены", "Выходные токены"]
                        ]
                    )

                with col_s2:
                    st.markdown("##### ⏱️ Динамика задержки ответа (в секундах)")
                    st.line_chart(df_m.set_index("Запрос №")[["Время ответа (сек)"]])

                st.markdown("##### 📜 Подробный журнал сессии")
                st.dataframe(
                    df_m,
                    column_config={
                        "Запрос №": st.column_config.NumberColumn("№", width="small"),
                        "Входные токены": st.column_config.NumberColumn(
                            "Входные (Prompt)", format="%d"
                        ),
                        "Выходные токены": st.column_config.NumberColumn(
                            "Выходные (Gen)", format="%d"
                        ),
                        "Всего токенов": st.column_config.NumberColumn(
                            "Всего токенов", format="%d"
                        ),
                        "Время ответа (сек)": st.column_config.NumberColumn(
                            "Время (сек)", format="%.2f s"
                        ),
                        "Поиск Qdrant (мс)": st.column_config.NumberColumn(
                            "Qdrant (мс)", format="%d ms"
                        ),
                        "Проект": "Проект",
                    },
                    use_container_width=True,
                    hide_index=True,
                )

    logs_data = get_audit_logs()
    df_logs_all = pd.DataFrame(logs_data) if logs_data else pd.DataFrame()

    if "📜 Полный Журнал Логов" in s_tab_dict:
        with s_tab_dict["📜 Полный Журнал Логов"]:
            st.write("История всех действий фиксируется в Qdrant Cloud:")
            if df_logs_all.empty:
                st.info("Журнал аудита пуст.")
            else:
                st.dataframe(
                    df_logs_all[[
                        "timestamp",
                        "username",
                        "role",
                        "ip",
                        "country",
                        "city",
                        "action",
                        "details",
                    ]],
                    use_container_width=True,
                )

    if "🗺️ Карта Входов (GeoIP)" in s_tab_dict:
        with s_tab_dict["🗺️ Карта Входов (GeoIP)"]:
            st.markdown("### 🗺️ Интерактивная карта геопозиций входов")
            if (
                not df_logs_all.empty
                and "lat" in df_logs_all.columns
                and "lon" in df_logs_all.columns
            ):
                df_map_data = df_logs_all.dropna(subset=["lat", "lon"]).copy()
                df_map_data["lat"] = pd.to_numeric(df_map_data["lat"], errors="coerce")
                df_map_data["lon"] = pd.to_numeric(df_map_data["lon"], errors="coerce")
                df_map_clean = df_map_data.dropna(subset=["lat", "lon"])

                if not df_map_clean.empty:
                    st.map(df_map_clean[["lat", "lon"]], zoom=2)
                    st.divider()
                    st.markdown("### 🌐 Распределение входов по странам и городам")
                    geo_summary = (
                        df_map_clean.groupby(["country", "city", "ip"])
                        .size()
                        .reset_index(name="Подключений")
                    )
                    st.dataframe(geo_summary, use_container_width=True)
                else:
                    st.info("Координаты подключений пока не зафиксированы.")
            else:
                st.info("Нет данных для отображения карты.")

    if "💡 Пробелы в знаниях & Отзывы" in s_tab_dict:
        with s_tab_dict["💡 Пробелы в знаниях & Отзывы"]:
            st.markdown(
                "### 🔍 1. Вопросы, на которые AI не нашел ответа (Knowledge Gaps)"
            )
            if not df_logs_all.empty and "action" in df_logs_all.columns:
                df_gaps = df_logs_all[df_logs_all["action"] == "KNOWLEDGE_GAP"]
                if df_gaps.empty:
                    st.success("🎉 Вопросов без ответа не зафиксировано.")
                else:
                    st.dataframe(
                        df_gaps[["timestamp", "username", "ip", "details"]],
                        use_container_width=True,
                    )
            else:
                st.info("Данные отсутствуют.")

            st.divider()
            st.markdown("### 👎 2. Замечания и негативные отзывы пользователей")
            if not df_logs_all.empty and "action" in df_logs_all.columns:
                df_neg = df_logs_all[df_logs_all["action"] == "FEEDBACK_NEGATIVE"]
                if df_neg.empty:
                    st.success("🎉 Замечаний от пользователей пока нет.")
                else:
                    st.dataframe(
                        df_neg[["timestamp", "username", "ip", "details"]],
                        use_container_width=True,
                    )
            else:
                st.info("Замечания отсутствуют.")

    if "👥 Управление Аккаунтами" in s_tab_dict:
        with s_tab_dict["👥 Управление Аккаунтами"]:
            st.markdown("### 👥 Список зарегистрированных пользователей")
            for login_key, u_info in st.session_state.users_db.items():
                with st.expander(
                    f"👤 **{u_info['name']}** (`{login_key}`) — Роль:"
                    f" `{role_badges.get(u_info['role'], u_info['role'])}`",
                    expanded=True,
                ):
                    col_u1, col_u2, col_u3 = st.columns([2, 2, 2])
                    with col_u1:
                        is_blk = u_info.get("is_blocked", False)
                        st.write(
                            f"**Статус:** {'🔴 ЗАБЛОКИРОВАН' if is_blk else '🟢 Активен'}"
                        )
                    with col_u2:
                        st.write(
                            f"**Лимит сессий:** `{u_info.get('max_connections', 1)}`"
                        )
                    with col_u3:
                        if is_blk:
                            if st.button("🔓 Разблокировать", key=f"unblk_{login_key}"):
                                u_info["is_blocked"] = False
                                u_info["failed_attempts"] = 0
                                st.success("Пользователь разблокирован!")
                                st.rerun()
